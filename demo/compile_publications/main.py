import json
from docx import Document

# 记录所有条目至一个列表中
items = []
for name in ['name1', 'name2']:
    fp_r = open(f'{name}.json', 'r')
    items += json.load(fp_r)
    fp_r.close()
print(f'合并条目数量：{len(items)}')

new_items = dict()
document = Document()


def parse_authors(authors):
    au_str = ''
    if len(authors) == 1:
        author = authors[0]
        au_str = f"{author['given']} {author['family']},"
    else:
        for author in authors[:-1]:
            au_str += f"{author['given']} {author['family']}, "
        author = authors[-1]
        au_str += f"and {author['given']} {author['family']},"
    return au_str


def parse_journal(journal):
    abbrs = {
        'Inf. Fusion': 'Information Fusion',
        'Trans. Pattern Anal. Mach. Intell.':
        'Transactions on Pattern Analysis and Machine Intelligence',
        'Signal Process. Lett.': 'Signal Processing Letters',
        'Trans. Image Process.': 'Transactions on Image Processing',
        'Trans. Multim.': 'Transactions on Multimedia',
        'Trans. Medical Imaging': 'Transactions on Medical Imaging',
        'Geosci. Remote. Sens. Lett.': 'Geoscience and Remote Sensing Letters',
        'Int. J. Comput. Vis.': 'International Journal of Computer Vision',
        'J. Biomed. Health Informatics':
        'Journal of Biomedical and Health Informatics',
        'Trans. Circuits Syst. Video Technol.':
        'Transactions on Circuits and Systems for Video Technology',
        'Internet Things J.': 'Internet of Things Journal',
        'Trans. Syst. Man Cybern. Syst.':
        'Transactions on Systems, Man, and Cybernetics',
        'J. Sel. Areas Commun.': 'Journal on Selected Areas in Communications',
        'J. Sel. Top. Signal Process.':
        'Journal of Selected Topics in Signal Processing',
        'Pattern Recognit.': 'Pattern Recognition',
        'Trans. Broadcast.': 'Transactions on Broadcasting',
        'Trans. Commun.': 'Transactions on Communications',
        'J. Vis. Commun. Image Represent.':
        'Journal of Visual Communication and Image Representation',
        'Frontiers Inf. Technol. Electron. Eng.':
        'Frontiers of Information Technology & Electronic Engineering',
        'Trans. Geosci. Remote. Sens.':
        'Transactions on Geoscience and Remote Sensing',
        'Comput. Vis. Image Underst.':
        'Computer Vision and Image Understanding',
        'J. Adv. Signal Process.': 'Journal on Advances in Signal Processing',
        'IET Commun.': 'IET Communications',
        'Signal Process. Image Commun.':
        'Signal Processing: Image Communication',
        'Pattern Recognition Lett.': 'Pattern Recognition Letters',
        'Wirel. Pers. Commun.': 'Wireless Personal Communications',
        'J. Wirel. Commun. Netw.':
        'Journal on Wireless Communications and Networking',
        'Trans. Veh. Technol.': 'Transactions on Vehicular Technology',
        'Int. J. Pattern Recognition Artif. Intell.':
        'International Journal of Pattern Recognition and Artificial Intelligence',
        'Commun. Lett.': 'Communications Letters',
        'Wirel. Communications Letters': 'Wireless Communications Letters',
        'Trans. Inf. Theory': 'Transactions on Information Theory',
        'Wirel. Commun. Mob. Comput.':
        'Wireless Communications and Mobile Computing',
        'Trans. Emerg. Telecommun. Technol.':
        'Transactions on Emerging Telecommunications Technologies',
        'Trans. Dependable Secur. Comput.':
        'Transactions on Dependable and Secure Computing',
        'J. Syst. Archit.': 'Journal of Systems Architecture',
        'Trans. Circuits Syst. II Express Briefs':
        'Transactions on Circuits and Systems II: Express Briefs',
        'Trans. Serv. Comput.': 'Transactions on Services Computing',
        'Comput. Secur.': 'Computers & Security',
        'Concurr. Comput. Pract. Exp.':
        'Concurrency and Computation: Practice and Experience',
        'Frontiers Comput. Sci.': 'Frontiers of Computer Science',
        'Inf. Sci.': 'Information Sciences',
        'Trans. Inf. Forensics Secur.':
        'Transactions on Information Forensics and Security',
        'Comput. Commun.': 'Computer Communications',
        'J. Commun. Inf. Networks':
        'Journal of Communications and Information Networks',
        'J. Comput. Sci. Technol.':
        'Journal of Computer Science and Technology',
        'Trans. Inf. Syst.': 'Transactions on Information and Systems',
        'J. Electronic Imaging': 'Journal of Electronic Imaging',
        'Electron. Express': 'Electronics Express',
        'Des. Test': 'Design & Test of Computers',
        'Trans. Computational Imaging':
        'Transactions on Computational Imaging',
        'Expert Syst. Appl.': 'Expert Systems with Applications',
    }
    for abbr in abbrs:
        if abbr in journal:
            journal = journal.replace(abbr, abbrs[abbr])
    return journal


def parse_conference(conference):
    knowns = {
        'IEEE intl conf on parallel & distributed processing with applications':
        'IEEE intl conf on parallel & distributed processing with applications, big data & cloud computing, sustainable computing & communications, social computing & networking',
        'IEEE international conference on acoustics':
        'IEEE international conference on acoustics, speech and signal processing',
        'international conference on image processing theory':
        'international conference on image processing theory, tools and applications',
        'IEEE international conference on electronics':
        'IEEE international conference on electronics, circuits, and systems',
        'IEEE intl conference on big data security on cloud':
        'IEEE intl conference on big data security on cloud, IEEE intl conference on high performance and smart computing and IEEE intl conference on intelligent data and security',
        'Knowledge science':
        'Knowledge science, engineering and management',
        'IEEE international conference on trust':
        'IEEE international conference on trust, security and privacy in computing and communications',
    }
    abbrs = {
        'intl conf': 'international conference',
    }

    conf_str = conference.split(',')[0]
    for known in knowns:
        if known in conf_str:
            conf_str = conf_str.replace(known, knowns[known])
            break
    for abbr in abbrs:
        if abbr in conf_str:
            conf_str = conf_str.replace(abbr, abbrs[abbr])
    return conf_str


for item in items:
    # 为了方便最后按年份排序，把年份放 ID 前面
    year = item['issued']['date-parts'][0][0]
    id = str(year) + '-' + item['id']

    # 跳过重复的
    if id in new_items.keys():
        continue

    if item['type'] in [
            'article-journal',
            'paper-conference',
    ]:
        au_str = parse_authors(item['author'])
        title = item['title']
        container = item['container-title']

        new_items[id] = [f"{au_str} \"{title},\""]  # 第 1/3 项

        if item['type'] == 'article-journal':  # 期刊
            container = parse_journal(container)  # 修改有问题的期刊名
        elif item['type'] == 'paper-conference':  # 会议
            container = parse_conference(container)  # 修改有问题的会议名
        new_items[id].append(f' {container}')  # 第 2/3 项；需要斜体，单独列出

        _str = ''
        if 'volume' in item:
            _str += f", vol. {item['volume']}"
        if 'number' in item:
            _str += f", no. {item['number']}"
        if 'page' in item:
            _str += f", pp. {item['page']}"
        _str += f', {year}.'
        new_items[id].append(_str)  # 第 3/3 项

print(f'处理后条目数量：{len(new_items)}')

# 排序，加序号，打到 word 里

sorted_ids = sorted(new_items.keys(), reverse=True)
idx = 0
for id in sorted_ids:
    idx += 1
    p = document.add_paragraph(f'[{idx}] ')
    p.add_run(new_items[id][0])
    p.add_run(new_items[id][1]).italic = True
    p.add_run(new_items[id][2])

document.save('output.docx')
